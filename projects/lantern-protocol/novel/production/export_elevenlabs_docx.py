#!/usr/bin/env python3
"""Build an ElevenLabs-ready audiobook manuscript DOCX.

The exporter intentionally emits reader-facing manuscript text only. It removes
chapter planning metadata such as canon sources, POV strategy, continuity notes,
and internal assembly comments while preserving in-world document inserts that
belong in the narration.
"""
from __future__ import annotations

import argparse
import re
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, List, Sequence

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - exercised by user environment
    raise SystemExit(
        "python-docx is required for ElevenLabs DOCX export. "
        "Install it with: python -m pip install python-docx"
    ) from exc

ROOT = Path(__file__).resolve().parents[1]
FRONT_MATTER = ROOT / "manuscript" / "front-matter"
CHAPTERS = ROOT / "manuscript" / "chapters"
EXPORTS = ROOT / "exports"

DEFAULT_OUTPUT_DOCX = EXPORTS / "Lantern_Protocol_Book_One_ElevenLabs.docx"
DEFAULT_REPORT = EXPORTS / "lantern-protocol-elevenlabs-audiobook-report.md"
MANUSCRIPT_START = "## Manuscript"
NOTES_STARTERS = ["## Continuity Notes", "## Revision Notes"]
INTERNAL_SECTION_HEADINGS = {
    "Canon Sources",
    "POV Strategy",
    "Chapter Purpose",
    "Continuity Requirements",
    "Continuity Notes",
    "Revision Notes",
}
INTERNAL_MARKERS = [
    "Canon Sources",
    "POV Strategy",
    "Chapter Purpose",
    "Continuity Requirements",
    "<!-- SOURCE:",
    "Auto-assembled",
]
EXPECTED_ACTIVE_CHAPTERS = 24
DETERMINISTIC_DOCX_TIMESTAMP = (2026, 1, 1, 0, 0, 0)
DETERMINISTIC_CORE_TIMESTAMP = datetime(2026, 1, 1, tzinfo=timezone.utc)


@dataclass(frozen=True)
class ExportResult:
    output_docx: Path
    report_path: Path
    front_matter_count: int
    chapter_count: int
    body_word_count: int
    findings: List[str]


def chapter_sort_key(path: Path) -> int:
    match = re.search(r"chapter-(\d+)", path.name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else 999


def front_matter_sort_key(path: Path) -> int:
    order = {
        "title": 1,
        "preface": 10,
        "foreword": 20,
        "introduction": 30,
        "manifest": 40,
        "author-note": 50,
    }
    return order.get(path.stem.lower(), 999)


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+(?:['-]\w+)?\b", text))


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def get_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    return fallback


def strip_markdown_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^[-*+]\s+", "", text)
    text = re.sub(r"^>\s?", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


def remove_internal_sections(text: str) -> str:
    """Remove known planning-only Markdown sections from front matter."""
    lines = text.splitlines()
    kept: List[str] = []
    skipping = False
    for line in lines:
        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if heading:
            heading_text = heading.group(2).strip().strip("#").strip()
            skipping = heading_text in INTERNAL_SECTION_HEADINGS
            if skipping:
                continue
        if not skipping:
            kept.append(line)
    return "\n".join(kept).strip()


def manuscript_body(text: str) -> str:
    start = text.find(MANUSCRIPT_START)
    if start == -1:
        body = remove_internal_sections(text)
    else:
        body = text[start + len(MANUSCRIPT_START):]
    earliest_note = len(body)
    for marker in NOTES_STARTERS:
        idx = body.find(marker)
        if idx != -1:
            earliest_note = min(earliest_note, idx)
    return body[:earliest_note].strip()


def reader_facing_front_matter(text: str) -> str:
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    text = remove_internal_sections(text)
    return text.strip()


def configure_document(document: DocumentType) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True


def add_bold_heading(document: DocumentType, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    run.bold = True


def add_title_page(document: DocumentType, book_title: str, book_number: str | None) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(book_title)
    run.bold = True

    if book_number:
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(book_number)
        subtitle_run.bold = True
        subtitle_run.font.size = Pt(14)

    document.add_page_break()


def add_insert_paragraph(document: DocumentType, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_body_paragraph(document: DocumentType, text: str) -> None:
    cleaned = strip_markdown_inline(text)
    if cleaned:
        document.add_paragraph(cleaned)


def add_markdown_text(document: DocumentType, text: str, *, allow_heading_one: bool = False) -> int:
    """Render a narrow manuscript-safe Markdown subset into the DOCX.

    Returns the reader-facing word count added to the document.
    """
    added_words = 0
    in_fence = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("<!--"):
            continue
        if stripped == "---":
            continue
        if stripped.startswith("```"):
            in_fence = not in_fence
            continue
        if in_fence:
            cleaned_insert = strip_markdown_inline(stripped)
            if cleaned_insert:
                add_insert_paragraph(document, cleaned_insert)
                added_words += word_count(cleaned_insert)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if heading:
            level = len(heading.group(1))
            heading_text = strip_markdown_inline(heading.group(2))
            if heading_text in INTERNAL_SECTION_HEADINGS:
                continue
            if level == 1 and not allow_heading_one:
                continue
            add_bold_heading(document, heading_text, level=1 if level == 1 else 2)
            added_words += word_count(heading_text)
            continue

        cleaned = strip_markdown_inline(stripped)
        if cleaned:
            add_body_paragraph(document, cleaned)
            added_words += word_count(cleaned)
    return added_words


def collect_front_matter() -> List[Path]:
    if not FRONT_MATTER.exists():
        return []
    return sorted(FRONT_MATTER.glob("*.md"), key=front_matter_sort_key)


def collect_chapters() -> List[Path]:
    return sorted(CHAPTERS.glob("chapter-*.md"), key=chapter_sort_key)


def build_report(result: ExportResult) -> str:
    status = "PASS" if not result.findings else "WARN"
    output_path = result.output_docx.relative_to(ROOT).as_posix()
    lines = [
        "# Lantern Protocol — ElevenLabs Audiobook Export Report",
        "",
        f"**Status:** {status}",
        f"**Output DOCX:** `{output_path}`",
        f"**Front matter files included:** {result.front_matter_count}",
        f"**Chapters included:** {result.chapter_count}",
        f"**Reader-facing words exported:** {result.body_word_count}",
        "",
        "## Validation Gates",
        "",
        "- Book title is emitted with the Word `Title` style and explicit bold run formatting.",
        "- Chapter headings are emitted with Word heading style and explicit bold run formatting.",
        "- Chapter bodies are taken only from `## Manuscript` through the next continuity/revision section boundary.",
        "- Internal planning sections are omitted from the audiobook export.",
        "- Markdown code fences are normalized into styled in-world insert paragraphs.",
        "",
        "## Findings",
        "",
    ]
    if result.findings:
        lines.extend(f"- {finding}" for finding in result.findings)
    else:
        lines.append("- No export warnings detected.")
    lines.append("")
    return "\n".join(lines)


def validate_export(front_matter_files: Sequence[Path], chapters: Sequence[Path], rendered_text: str) -> List[str]:
    findings: List[str] = []
    if not chapters:
        findings.append(f"Critical: no chapter files found under `{CHAPTERS}`.")
    if len(chapters) != EXPECTED_ACTIVE_CHAPTERS:
        findings.append(f"Expected {EXPECTED_ACTIVE_CHAPTERS} active chapters, found {len(chapters)}.")
    for chapter in chapters:
        body = manuscript_body(read_text(chapter))
        if not body.strip():
            findings.append(f"{chapter.name} has no reader-facing manuscript body.")
    for marker in INTERNAL_MARKERS:
        if marker.lower() in rendered_text.lower():
            findings.append(f"Internal marker leaked into export text: `{marker}`.")
    if not front_matter_files:
        findings.append("No front matter files were found; export contains title page and chapters only.")
    return findings


def normalize_docx_package(path: Path) -> None:
    """Rewrite DOCX zip metadata so CI output is stable across platforms."""
    tmp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path) as src, zipfile.ZipFile(tmp_path, "w") as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            normalized = zipfile.ZipInfo(item.filename, DETERMINISTIC_DOCX_TIMESTAMP)
            normalized.compress_type = item.compress_type
            normalized.external_attr = item.external_attr
            normalized.comment = item.comment
            normalized.create_system = item.create_system
            dst.writestr(normalized, data)
    tmp_path.replace(path)


def build_elevenlabs_docx(
    *,
    book_title: str = "Lantern Protocol",
    book_number: str | None = "Book One",
    output_docx: Path = DEFAULT_OUTPUT_DOCX,
    report_path: Path = DEFAULT_REPORT,
) -> ExportResult:
    EXPORTS.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    document.core_properties.created = DETERMINISTIC_CORE_TIMESTAMP
    document.core_properties.modified = DETERMINISTIC_CORE_TIMESTAMP
    add_title_page(document, book_title, book_number)

    front_matter_files = collect_front_matter()
    chapters = collect_chapters()
    exported_fragments: List[str] = [book_title, book_number or ""]
    exported_words = word_count(book_title) + word_count(book_number or "")

    for front_matter in front_matter_files:
        raw_text = read_text(front_matter)
        title = get_title(raw_text, front_matter.stem.replace("-", " ").title())
        clean_text = reader_facing_front_matter(raw_text)
        add_bold_heading(document, title, level=1)
        exported_words += word_count(title)
        exported_words += add_markdown_text(document, clean_text, allow_heading_one=False)
        exported_fragments.append(title)
        exported_fragments.append(clean_text)
        document.add_page_break()

    for idx, chapter in enumerate(chapters, start=1):
        raw_text = read_text(chapter)
        title = get_title(raw_text, chapter.stem.replace("-", " ").title())
        body = manuscript_body(raw_text)
        add_bold_heading(document, title, level=1)
        exported_words += word_count(title)
        exported_words += add_markdown_text(document, body, allow_heading_one=False)
        exported_fragments.append(title)
        exported_fragments.append(body)
        if idx < len(chapters):
            document.add_page_break()

    rendered_text = "\n".join(exported_fragments)
    findings = validate_export(front_matter_files, chapters, rendered_text)
    result = ExportResult(
        output_docx=output_docx,
        report_path=report_path,
        front_matter_count=len(front_matter_files),
        chapter_count=len(chapters),
        body_word_count=exported_words,
        findings=findings,
    )
    document.save(output_docx)
    normalize_docx_package(output_docx)
    report_path.write_text(build_report(result), encoding="utf-8")
    return result


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build an ElevenLabs-ready audiobook DOCX export.")
    parser.add_argument("--book-title", default="Lantern Protocol", help="Book title to place on the title page.")
    parser.add_argument("--book-number", default="Book One", help="Subtitle/book number. Pass an empty string to omit.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT_DOCX), help="DOCX output path.")
    parser.add_argument("--report", default=str(DEFAULT_REPORT), help="Markdown validation report path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    result = build_elevenlabs_docx(
        book_title=args.book_title,
        book_number=args.book_number or None,
        output_docx=Path(args.output),
        report_path=Path(args.report),
    )
    print(f"Wrote {result.output_docx}")
    print(f"Wrote {result.report_path}")
    print(f"Front matter files: {result.front_matter_count}")
    print(f"Total chapters: {result.chapter_count}")
    print(f"Reader-facing words: {result.body_word_count}")
    if result.findings:
        print("Warnings:")
        for finding in result.findings:
            print(f"- {finding}")


if __name__ == "__main__":
    main()
