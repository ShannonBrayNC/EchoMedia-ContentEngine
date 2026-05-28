#!/usr/bin/env python3
"""Generate Book 1 ElevenLabs-ready DOCX from canonical Lantern manuscript files.

Run from repo root:
    python tools/generate_book1_elevenlabs_docx.py

Output:
    dist/Lantern_Protocol_Book_1_The_Living_Anchor_ElevenLabs.docx
"""

from __future__ import annotations

import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: python-docx\n"
        "Install it with:\n"
        "    python -m pip install python-docx\n"
    ) from exc

ROOT = Path.cwd()
NOVEL = ROOT / "projects" / "lantern-protocol" / "novel"
MANUSCRIPT = NOVEL / "manuscript"
CHAPTERS = MANUSCRIPT / "chapters"
FRONT = MANUSCRIPT / "front-matter"
BACK = MANUSCRIPT / "back-matter"
DIST = ROOT / "dist"

TITLE_PAGE = FRONT / "title-page-elevenlabs.md"
PREFACE = FRONT / "preface.md"
BOOK2_TEASER = BACK / "book-2-teaser.md"
OUTPUT = DIST / "Lantern_Protocol_Book_1_The_Living_Anchor_ElevenLabs.docx"
CHAPTER_RE = re.compile(r"^chapter-(\d{2})-.+\.md$")

DROP_SECTIONS = {
    "Canon Sources",
    "POV Strategy",
    "Chapter Purpose",
    "Continuity Requirements",
    "Continuity Notes",
    "Revision Notes",
    "Processing Metadata",
}


def read(path: Path, required: bool = True) -> str:
    if not path.exists():
        if required:
            raise FileNotFoundError(f"Missing required file: {path}")
        return ""
    return path.read_text(encoding="utf-8").replace("\r\n", "\n").strip()


def extract_section(text: str, heading: str, end_pattern: str) -> str:
    start = text.find(heading)
    if start < 0:
        return ""
    rest = text[start + len(heading):]
    match = re.search(end_pattern, rest, flags=re.MULTILINE)
    return (rest[: match.start()] if match else rest).strip()


def extract_narration_script(text: str) -> str:
    return extract_section(text, "# Narration Script", r"\n##\s+") or text


def extract_manuscript(text: str) -> str:
    body = extract_section(text, "## Manuscript", r"\n##\s+(Continuity Notes|Revision Notes|Processing Metadata)\b")
    if not body:
        raise ValueError("Chapter missing ## Manuscript section")
    return body.strip()


def chapter_paths() -> list[Path]:
    if not CHAPTERS.exists():
        raise FileNotFoundError(f"Chapter folder not found: {CHAPTERS}")
    files = [p for p in CHAPTERS.iterdir() if CHAPTER_RE.match(p.name)]
    files = sorted(files, key=lambda p: int(CHAPTER_RE.match(p.name).group(1)))
    if len(files) != 24:
        raise SystemExit(f"Expected 24 chapter files, found {len(files)}")
    return files


def strip_fences(lines: list[str]) -> list[str]:
    cleaned = []
    in_fence = False
    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_fence = not in_fence
            continue
        cleaned.append(line)
    return cleaned


def add_inline_md(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_para(doc: Document, text: str, style: str | None = None, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_inline_md(p, text)
    if italic:
        for run in p.runs:
            run.italic = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    clean = text.strip().lstrip("#").strip()
    p = doc.add_heading(level=level)
    run = p.add_run(clean)
    run.bold = True


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)


def add_markdown(doc: Document, markdown: str, break_before: bool = False) -> None:
    if break_before:
        page_break(doc)
    for line in strip_fences(markdown.splitlines()):
        stripped = line.strip()
        if not stripped or stripped == "---" or stripped.startswith("<!-- SOURCE:"):
            continue
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
        elif stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading not in DROP_SECTIONS:
                add_heading(doc, heading, 2)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
        elif stripped.startswith("- "):
            add_para(doc, stripped[2:], "List Bullet")
        elif stripped.startswith("> "):
            add_para(doc, stripped[2:], italic=True)
        elif stripped.lower() == "[pause]":
            add_para(doc, "[pause]", italic=True)
        else:
            add_para(doc, stripped)


def build() -> Path:
    DIST.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lantern Protocol")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Book One: The Living Anchor")
    sub_run.bold = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("Written by Shannon Bray")

    add_markdown(doc, extract_narration_script(read(TITLE_PAGE)), break_before=True)

    preface = read(PREFACE, required=False)
    if preface:
        add_markdown(doc, preface, break_before=True)

    for path in chapter_paths():
        content = read(path)
        heading_match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        heading = heading_match.group(1).strip() if heading_match else path.stem
        body = extract_manuscript(content)
        page_break(doc)
        add_heading(doc, heading, 1)
        add_markdown(doc, body)

    teaser = read(BOOK2_TEASER, required=False)
    if teaser:
        add_markdown(doc, teaser, break_before=True)

    core = doc.core_properties
    core.author = "Shannon Bray"
    core.title = "Lantern Protocol: Book One - The Living Anchor"
    core.subject = "ElevenLabs narration manuscript"
    core.keywords = "Lantern Protocol, ElevenLabs, audiobook, narration manuscript"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build()
    print(f"Generated: {output}")
