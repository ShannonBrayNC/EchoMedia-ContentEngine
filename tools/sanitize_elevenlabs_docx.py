from __future__ import annotations

import re
from pathlib import Path
from docx import Document

ROOT = Path.cwd()

DOCX_FILES = [
    ROOT / "dist" / "elevenlabs" / "book-1" / "Lantern_Protocol_Book_1_The_Living_Anchor_ElevenLabs.docx",
    ROOT / "dist" / "elevenlabs" / "book-2" / "Lantern_Protocol_Book_2_The_Separate_Agreements_ElevenLabs.docx",
]

PAUSE_PATTERNS = [
    r"\[pause\]",
    r"\[long pause\]",
    r"\[short pause\]",
    r"\(pause\)",
    r"\(long pause\)",
    r"\(short pause\)",
]

def clean_text(value: str) -> str:
    text = value

    for pattern in PAUSE_PATTERNS:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)

    # ElevenLabs reads underscores awkwardly. Convert them to spaces.
    text = text.replace("_", " ")

    # Clean spacing caused by removed pause markers.
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s+([,.!?;:])", r"\1", text)

    return text.strip()

def clean_paragraph(paragraph) -> None:
    if not paragraph.runs:
        return

    original = "".join(run.text for run in paragraph.runs)
    cleaned = clean_text(original)

    if original == cleaned:
        return

    # Preserve the style of the first run. This intentionally flattens mixed-run
    # bold/italic inside affected paragraphs, which is acceptable for narration.
    first = paragraph.runs[0]
    for run in paragraph.runs:
        run.text = ""

    first.text = cleaned

def clean_docx(path: Path) -> None:
    if not path.exists():
        print(f"SKIP missing: {path}")
        return

    doc = Document(path)

    for paragraph in doc.paragraphs:
        clean_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    clean_paragraph(paragraph)

    doc.save(path)
    print(f"CLEANED: {path}")

def main() -> None:
    for path in DOCX_FILES:
        clean_docx(path)

if __name__ == "__main__":
    main()
