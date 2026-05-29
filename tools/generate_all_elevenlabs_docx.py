from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

ROOT = Path.cwd()

BOOKS = [
    {
        "key": "book-1",
        "title": "Lantern Protocol",
        "subtitle": "Book One: The Living Anchor",
        "author": "Shannon Bray",
        "chapters": ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "chapters",
        "front": ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "front-matter",
        "back": ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "back-matter",
        "output": ROOT / "dist" / "elevenlabs" / "book-1" / "Lantern_Protocol_Book_1_The_Living_Anchor_ElevenLabs.docx",
    },
    {
        "key": "book-2",
        "title": "Lantern Protocol II",
        "subtitle": "The Separate Agreements",
        "author": "Shannon Bray",
        "chapters": ROOT / "projects" / "lantern-protocol" / "novel" / "book-ii" / "manuscript" / "chapters",
        "front": ROOT / "projects" / "lantern-protocol" / "novel" / "book-ii" / "manuscript" / "front-matter",
        "back": ROOT / "projects" / "lantern-protocol" / "novel" / "book-ii" / "manuscript" / "back-matter",
        "output": ROOT / "dist" / "elevenlabs" / "book-2" / "Lantern_Protocol_Book_2_The_Separate_Agreements_ElevenLabs.docx",
    },
]

CHAPTER_RE = re.compile(r"^chapter-(\d{2})-.+\.md$")


def read_optional(path: Path) -> str:
    return path.read_text(encoding="utf-8").replace("\r\n", "\n").strip() if path.exists() else ""


def extract_section(content: str, start_heading: str, end_pattern: str) -> str:
    start = content.find(start_heading)
    if start < 0:
        return ""
    rest = content[start + len(start_heading):]
    match = re.search(end_pattern, rest, flags=re.MULTILINE)
    return (rest[: match.start()] if match else rest).strip()


def extract_manuscript(content: str) -> str:
    body = extract_section(
        content,
        "## Manuscript",
        r"\n##\s+(Continuity Notes|Revision Notes|Processing Metadata)\b",
    )
    return body.strip()


def extract_narration_script(content: str) -> str:
    return extract_section(content, "# Narration Script", r"\n##\s+") or content


def chapter_files(chapter_dir: Path) -> list[Path]:
    if not chapter_dir.exists():
        raise FileNotFoundError(f"Chapter folder not found: {chapter_dir}")

    files = [p for p in chapter_dir.iterdir() if CHAPTER_RE.match(p.name)]
    return sorted(files, key=lambda p: int(CHAPTER_RE.match(p.name).group(1)))


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    text = text.strip().lstrip("#").strip()
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.bold = True


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def add_para(doc: Document, text: str, style: str | None = None, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_inline(p, text)
    if italic:
        for run in p.runs:
            run.italic = True


def add_markdown(doc: Document, markdown: str, break_before: bool = False) -> None:
    if not markdown.strip():
        return

    if break_before:
        page_break(doc)

    in_code = False

    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if not stripped:
            continue

        if stripped == "---":
            continue

        if stripped.startswith("<!-- SOURCE:"):
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:], 1)
            continue

        if stripped.startswith("## "):
            # Skip internal metadata headings when source sections leak in.
            heading = stripped[3:].strip()
            if heading in {
                "Canon Sources",
                "POV Strategy",
                "Chapter Purpose",
                "Continuity Requirements",
                "Continuity Notes",
                "Revision Notes",
                "Processing Metadata",
            }:
                continue
            add_heading(doc, heading, 2)
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 3)
            continue

        if stripped.startswith("- "):
            add_para(doc, stripped[2:], style="List Bullet")
            continue

        if stripped.startswith("> "):
            add_para(doc, stripped[2:], italic=True)
            continue

        if stripped.lower() == "[pause]":
            add_para(doc, "[pause]", italic=True)
            continue

        add_para(doc, stripped)


def build_book(config: dict) -> Path:
    output = config["output"]
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_doc(doc)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(config["title"])
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(config["subtitle"])
    subtitle_run.bold = True

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run(f"Written by {config['author']}")

    title_page = read_optional(config["front"] / "title-page-elevenlabs.md")
    if title_page:
        add_markdown(doc, extract_narration_script(title_page), break_before=True)

    preface = read_optional(config["front"] / "preface.md")
    if preface:
        add_markdown(doc, preface, break_before=True)

    files = chapter_files(config["chapters"])
    if not files:
        raise RuntimeError(f"No chapter files found under {config['chapters']}")

    for chapter in files:
        content = read_optional(chapter)
        heading_match = re.search(r"^#\s+(.+)$", content, flags=re.MULTILINE)
        heading = heading_match.group(1).strip() if heading_match else chapter.stem
        body = extract_manuscript(content)

        if not body:
            print(f"WARNING: {chapter.name} has no ## Manuscript body. Skipped.")
            continue

        page_break(doc)
        add_heading(doc, heading, 1)
        add_markdown(doc, body)

    for back_name in ("book-2-teaser.md", "book-3-teaser.md", "afterword.md"):
        back = read_optional(config["back"] / back_name)
        if back:
            add_markdown(doc, back, break_before=True)

    core = doc.core_properties
    core.author = config["author"]
    core.title = f"{config['title']}: {config['subtitle']}"
    core.subject = "ElevenLabs narration manuscript"
    core.keywords = "Lantern Protocol, ElevenLabs, audiobook, narration"

    doc.save(output)
    return output


def main() -> None:
    made = []
    for config in BOOKS:
        try:
            made.append(build_book(config))
        except Exception as exc:
            print(f"FAILED {config['key']}: {exc}")

    print()
    print("Generated files:")
    for path in made:
        print(f"  {path}")


if __name__ == "__main__":
    main()
