from __future__ import annotations

import re
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt

ROOT = Path.cwd()

BOOKS = [
    {
        "key": "book-1",
        "book_title": "The Living Anchor",
        "series": "Lantern Protocol",
        "subtitle": "Book One of the Lantern Protocol",
        "author": "Shannon Bray",
        "chapters": ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "chapters",
        "preface": ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "front-matter" / "preface.md",
        "back_matter": [
            ROOT / "projects" / "lantern-protocol" / "novel" / "manuscript" / "back-matter" / "book-2-teaser.md",
        ],
        "output": ROOT / "dist" / "kdp" / "book-1" / "The_Living_Anchor_KDP_6x9.docx",
    },
    {
        "key": "book-2",
        "book_title": "The Separate Agreements",
        "series": "Lantern Protocol",
        "subtitle": "Book Two of the Lantern Protocol",
        "author": "Shannon Bray",
        "chapters": ROOT / "projects" / "lantern-protocol" / "novel" / "book-ii" / "manuscript" / "chapters",
        "preface": ROOT / "projects" / "lantern-protocol" / "novel" / "book-ii" / "manuscript" / "front-matter" / "preface.md",
        "back_matter": [],
        "output": ROOT / "dist" / "kdp" / "book-2" / "The_Separate_Agreements_KDP_6x9.docx",
    },
]

CHAPTER_RE = re.compile(r"^chapter-(\d{2})-.+\.md$")


def read_optional(path: Path) -> str:
    return path.read_text(encoding="utf-8").replace("\r\n", "\n").strip() if path.exists() else ""


def extract_section(content: str, start_heading: str, end_pattern: str) -> str:
    start = content.find(start_heading)
    if start < 0:
        return ""
    rest = content[start + len(start_heading):]
    match = re.search(end_pattern, rest, flags=re.MULTILINE)
    return (rest[: match.start()] if match else rest).strip()


def extract_manuscript(content: str) -> str:
    return extract_section(
        content,
        "## Manuscript",
        r"\n##\s+(Continuity Notes|Revision Notes|Processing Metadata)\b",
    ).strip()


def chapter_files(chapters_dir: Path) -> list[Path]:
    if not chapters_dir.exists():
        raise FileNotFoundError(f"Chapter folder not found: {chapters_dir}")

    files = [p for p in chapters_dir.iterdir() if CHAPTER_RE.match(p.name)]
    return sorted(files, key=lambda p: int(CHAPTER_RE.match(p.name).group(1)))


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]

    # KDP-friendly common trim size: 6x9.
    section.page_width = Inches(6)
    section.page_height = Inches(9)

    # Comfortable print margins. Increase inside margin slightly for binding.
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.65)
    section.gutter = Inches(0.15)

    normal = doc.styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(11)
    normal.paragraph_format.first_line_indent = Inches(0.22)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Garamond"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(10)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int = 1, centered: bool = False) -> None:
    text = text.strip().lstrip("#").strip()
    p = doc.add_heading(level=level)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        else:
            paragraph.add_run(part)


def add_para(doc: Document, text: str, style: str | None = None, italic: bool = False, centered: bool = False, no_indent: bool = False) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if no_indent:
        p.paragraph_format.first_line_indent = Inches(0)
    add_inline(p, text)
    if italic:
        for run in p.runs:
            run.italic = True


def clean_story_line(line: str) -> str | None:
    stripped = line.strip()

    if not stripped:
        return None

    if stripped == "---":
        return None

    if stripped.startswith("<!-- SOURCE:"):
        return None

    if stripped.startswith("```"):
        return None

    if stripped.lower() == "[pause]":
        return None

    # Remove markdown table separator lines.
    if re.fullmatch(r"[\|\-\:\s]+", stripped):
        return None

    return stripped


def add_markdown_story(doc: Document, markdown: str, allow_h1: bool = True) -> None:
    in_code = False

    for raw in markdown.splitlines():
        stripped = raw.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        line = clean_story_line(raw)
        if line is None:
            continue

        if line.startswith("# "):
            if allow_h1:
                page_break(doc)
                add_heading(doc, line[2:], 1, centered=True)
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            if heading in {
                "Canon Sources",
                "POV Strategy",
                "Purpose",
                "Chapter Purpose",
                "Chapter Promise",
                "Continuity Requirements",
                "Continuity Notes",
                "Revision Notes",
                "Processing Metadata",
            }:
                continue
            add_heading(doc, heading, 2)
            continue

        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
            continue

        if line.startswith("- "):
            add_para(doc, line[2:], style="List Bullet", no_indent=True)
            continue

        if line.startswith("> "):
            add_para(doc, line[2:], italic=True)
            continue

        # KDP can handle story artifacts as ordinary paragraphs.
        # Avoid monospace/code styling for print simplicity.
        add_para(doc, line)


def add_title_page(doc: Document, book: dict) -> None:
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(book["book_title"])
    title_run.bold = True

    add_para(doc, book["subtitle"], centered=True, no_indent=True)
    add_para(doc, book["series"], centered=True, no_indent=True)
    add_para(doc, "", no_indent=True)
    add_para(doc, f"Written by {book['author']}", centered=True, no_indent=True)


def add_copyright_page(doc: Document, book: dict) -> None:
    page_break(doc)
    add_para(doc, f"{book['book_title']}", centered=True, no_indent=True)
    add_para(doc, f"{book['subtitle']}", centered=True, no_indent=True)
    add_para(doc, "", no_indent=True)
    add_para(doc, f"Copyright © 2026 {book['author']}", no_indent=True)
    add_para(doc, "All rights reserved.", no_indent=True)
    add_para(doc, "", no_indent=True)
    add_para(
        doc,
        "This is a work of fiction. Names, characters, organizations, technologies, places, events, and incidents are either products of the author's imagination or used fictitiously.",
        no_indent=True,
    )
    add_para(doc, "", no_indent=True)
    add_para(doc, "First edition.", no_indent=True)


def build_book(book: dict) -> tuple[Path, int, list[str]]:
    output = book["output"]
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_doc(doc)

    add_title_page(doc, book)
    add_copyright_page(doc, book)

    preface = read_optional(book["preface"])
    if preface:
        page_break(doc)
        add_markdown_story(doc, preface, allow_h1=True)

    warnings = []
    total_words = 0

    files = chapter_files(book["chapters"])
    if not files:
        raise RuntimeError(f"No chapter files found for {book['key']}")

    for path in files:
        content = read_optional(path)
        heading_match = re.search(r"^#\s+(.+)$", content, flags=re.MULTILINE)
        heading = heading_match.group(1).strip() if heading_match else path.stem
        body = extract_manuscript(content)

        words = len(re.findall(r"\b[\w'’-]+\b", body))
        total_words += words

        if words < 1000:
            warnings.append(f"{path.name}: small body ({words} words)")
        if re.search(r"Draft placeholder|Replace this section|TODO|TBD", content, flags=re.IGNORECASE):
            warnings.append(f"{path.name}: placeholder text detected")

        if not body:
            warnings.append(f"{path.name}: no manuscript body detected")
            continue

        page_break(doc)
        add_heading(doc, heading, level=1, centered=True)
        add_markdown_story(doc, body, allow_h1=False)

    for back_path in book["back_matter"]:
        back = read_optional(back_path)
        if back:
            page_break(doc)
            add_markdown_story(doc, back, allow_h1=True)

    core = doc.core_properties
    core.author = book["author"]
    core.title = book["book_title"]
    core.subject = book["subtitle"]
    core.keywords = f"{book['series']}, KDP, manuscript, fiction"

    doc.save(output)
    return output, total_words, warnings


def main() -> None:
    print("Generating KDP manuscripts...")
    print()

    for book in BOOKS:
        try:
            output, total_words, warnings = build_book(book)
            print(f"{book['key']}: {output}")
            print(f"  manuscript body words: {total_words}")

            if warnings:
                print("  WARNINGS:")
                for warning in warnings:
                    print(f"    - {warning}")
            else:
                print("  warnings: none")

            print()
        except Exception as exc:
            print(f"{book['key']}: FAILED")
            print(f"  {exc}")
            print()

    print("Done.")


if __name__ == "__main__":
    main()
